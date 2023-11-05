# headers 
__author__ = [
    "Alex Leedom <amleedom@gmail.com>"
]

__license__ = [
    "MIT License"
]

__version__ = [
    '0.2.0'
]

# import modules 
import os
import re
import pandas as pd
import requests 
import lxml.etree as ET 
from bs4 import BeautifulSoup
from collections import Counter 
from tqdm import tqdm

import sys
sys.path.append('~/cltk/open_words/')
import open_words.open_words.parse # seems to work 

from nltk.tokenize.punkt import PunktLanguageVars
# need to edit this path 
from cltk.lemmatize import LatinBackoffLemmatizer
#from cltk.corpus.utils.importer import CorpusImporter
from cltk.data.fetch import FetchCorpus

# following line deprecated by changes in the CLTK file structure 
# from cltk.corpus.readers import get_corpus_reader

# following line deprecated by changes in the CLTK file structure 
# from cltk.stop.latin import CorpusStoplist
from cltk.stops.lat import STOPS 
#from cltk.stop.latin import STOPS_LIST 

# need to import a corpus reader as well 
from cltkreaders.lat import LatinLibraryCorpusReader
from cltkreaders.lat import LatinTesseraeCorpusReader 

STOPS_LIST = STOPS + ['', ',','.','punc' 'ita', 'non', 'et', 'in', 'ab', 'a', 'e', 'p', 'm','ap', 'ego', 'sum', 'tu', 
                           'i','v','vi','vii','x','ix','I','V','X'] 

from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

lemmatizer = LatinBackoffLemmatizer()
path = 'lexica/Lewis_Short_XML/lat.ls.perseus-eng1.xml'
tree = ET.parse(path)
entries = tree.xpath('//entryFree')

punc =  [',', ';', '"', "'", '-que', '-ne', '-ve', 'punc', '?']

line_break = '———————————————————————————————————————'

# Eliminating the overall catalog for the time being in favor of one specific to LL 
latin_library_catalog = pd.read_csv('data/latin_library_catalog.csv')

tesserae_catalog = pd.read_csv('data/tesserae_catalog.csv')



parser = open_words.open_words.parse.Parse()

    # helper functions
    
def divide_chunks(l=[], n = 125):
    """
    Creates a generator for paragraphs in a text. Used to print to pages and define a small vocabulary chunk to perform lookups
    :param l: list to divide
    :param n: size of the chunk 
    """
    if l == []:
        l = list(self.paras)
    for i in range(0, len(l), n-1): 
        yield l[ i: i+n ]

def depth(l):
    """
    Helper function for flattening nested lists. Latin Library texts often come as nested lists which need to be flattened before parsing and formatting.
    """
    if isinstance(l, list):
        return 1 + max(depth(item) for item in l)
    else:
        return 0

def flatten_paragraphs(paras): 
    """
    Implements the depth function to flatten a list
    """
    while depth(paras) >= 2: 
        paras = [item for sublist in paras for item in sublist]
    return paras

def create_page(text_chunk, vocabulary_list = None, doc = None, save = False, title = ''): 
    """Creates page of text and vocabulary. 
    :param text_chunk: chunk of text to add to page
    """
    chunk_dict = {}
    br = line_break
    if doc == None: 
        doc = Document()
        
    main_paragraph = doc.add_paragraph().add_run(text_chunk)
    main_paragraph.font.size = Pt(12)
    # visual line break 
    p2 = doc.add_paragraph()
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(br)
    
    # column format 
    
    doc.add_section(0)
    section = doc.sections[-1]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    
    # add words 
    
    words = [w.replace('\n',' ').replace('—', ' ').strip('.,;:-—"').lower() for w in str(text_chunk).replace('\n', ' ').replace('—', ' ').split(' ')]

    words = [w.strip(' !,.;:"-—') for w in words]
    for word in set(words):
        chunk_dict[word] = {}
        
        # parse the nominative form. works better than the OW parser 
        if word == '':
            continue
        else:
            lemma = lemmatizer.lemmatize([word.lower()])[0][1]
            if lemma in vocabulary_list: 
#                 print(f'looking up {lemma}...')
                forms = lookup_word([lemma]).get(lemma)
        
                chunk_dict[word]['lemma'] = lemma
                
                chunk_dict[word]['dictionary_form'] = '' if forms.get('dictionary_form') is None\
                else forms.get('dictionary_form')
                
                chunk_dict[word]['senses'] = '' if forms.get('senses') is None\
                else forms.get('senses')
            
    # clean dictionary
    
    delete = [k for k in chunk_dict.keys() if chunk_dict[k].get('dictionary_form') == '']

    for k in delete:
        del chunk_dict[k]

    chunk_vocab = sorted(list(chunk_dict.keys()))
    # print(f"looked up {len(chunk_vocab)} words") # removed this because it was cluttering output window 
    if '' in chunk_vocab:
        chunk_vocab.remove('')
    
    # look up, format words 
    for vocab_word in chunk_vocab: 
        if chunk_dict[vocab_word] == {}:
            pass
        else:
            p = doc.add_paragraph()

            bold_text = f'{chunk_dict[vocab_word]["lemma"]} {" ".join(chunk_dict[vocab_word]["dictionary_form"]).strip(".")}: '
            plain_text = f'{chunk_dict[vocab_word]["senses"]}'

            bold_run = p.add_run()
            bold_run.bold = True
            bold_run.text = bold_text
            bold_run.font.size = Pt(10)

            plain_run = p.add_run()
            plain_run.bold = False
            plain_run.text = plain_text 
            plain_run.font.size = Pt(10)
            
            p.paragraph_format.space_after = Pt(0)

    doc.add_section(0)
    section = doc.sections[-1]

    # reset formatting

    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')
    doc.add_page_break()

    sections = doc.sections 
    for section in sections: 
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(.5)
        section.right_margin = Inches(.5)
    if save == True:
        doc.save(f'{title} with facing vocab.docx')  
    
def lookup_word(in_list = []): 
    passage_dict = {}
    for w in in_list: 

        endings = []
        gender = []

        # create dictionary for entries on the page 

        lemma = lemmatizer.lemmatize([w])[0][1]
        parsed = parser.parse(w)
        word = parsed['word']
        # word 
        passage_dict[word] = {}
        passage_dict[word]['word'] = word
        passage_dict[word]['lemma'] = lemma
        if len(parsed['defs']) >=1:
            passage_dict[word]['senses'] = [parsed['defs'][i]['senses'] for i in range(len(parsed['defs']))]
            passage_dict[word]['senses'] = ', '.join([meaning for sense in passage_dict[word]['senses'] for meaning in sense])
        else:
            continue
        # inflection
        for entry in entries:
            if entry.get('key') == lemma:
                if entry.find('itype') is not None: 
                    endings.append(f'{entry.find("itype").text} ')
                if entry.find('gen') is not None: 
                    endings.append(entry.find('gen').text)
        passage_dict[word]['dictionary_form'] = endings

    return passage_dict

class PharrBuilder():
    """
    This class makes lexical analyses of Latin texts available the Latin Library possible, and also allows the creation of DOCX documents with Latin text and facing vocabulary. 
    Currently, only Latin Library texts are available. 
    """
    def __init__(self, corpus_name: 'latin_library, tesserae'=''): 
        """
        For this version, we support tesserae and the latin library texts. 
        Call the catalog and select a text from the 'source' column. Note that this requires pandas syntax 

        Example: 

        cat = PharrBuilder(corpus_name = 'tesserae').catalog
        # choose a text 

        catilinam = 'cicero.in_catilinam.tess'
        """
        
        if corpus_name == 'latin_library':
            self.catalog = latin_library_catalog
            self.corpus_name = corpus_name
            # TODO instantiate Latin Library reader 
        elif corpus_name == 'tesserae':
            self.catalog =  tesserae_catalog = pd.read_csv('data/tesserae_catalog.csv')
            self.reader = LatinTesseraeCorpusReader()
            print(f'your reader is {self.reader}')
            self.corpus_name = corpus_name
        else:
            raise Exception('Invalid corpus selection. corpus_name input should be one of "latin_library" or "tesserae"')
        #self.reader = get_corpus_reader(language='lat')
        # TODO 
        # instantiate either perseus or latin library corpus
        
    
    # following lines deprecated since we select corpus at instantiation  
    # @classmethod 
    # def choose_corpus(self, corpus_name: 'perseus, latin_library, tesserae'='') -> "list": 
    #     """Lists the available Latin texts. Currently supports Latin Library and Perseus Library. Will display a list of available texts as a pandas `series`. 
    #     :Param corpus_name: currently we support perseus, the latin library, and tesserae 
    #     These are listed in the corpus_names attribute.  
    #     """
    #     if corpus_name == 'latin_library':
    #         self.catalog = latin_library_catalog # only LL catalog right now 
    #     if corpus_name == 'tesserae':
    #         self.catalog = tesserae_catalog

    #     #self.reader = FetchCorpus(language='latin', corpus_name=corpus_name)
    #     #self.corpus_name = corpus_name

    # @classmethod
    def get_texts(self, text_name: "source name", exclude_count: 'word count to exclude' = 10): 
        """Selects and parses a text pased by user. 
        
        Specify the text with the 'source' column of the catalog.
        :param text_name: retrieved from the `source` column of the catalog. Can be entered as either string or list. 
        :param exclude_count: when the word count of the text is calculated, we provide the dictionary form of only words that appear *fewer* than this number of times. The default is 10.
        """
        # get lemmas 
        corpus_name = self.corpus_name 
        if corpus_name == 'tesserae':
            # we have the tesserae reader and use its built-in lemmatizer to parse the words 
            
            lemmas = []
            reader = self.reader 
            text_words = reader.tokens(text_name)
            print(f'Compiling vocabulary in {text_name}')
            for token in tqdm(text_words):
                lemmas.append(token.lemma_) # don't want the actual words, just the lemmas will do 
            lemmas = [str(lemma) for lemma in lemmas if str(lemma) not in punc]
            
            word_count = Counter(lemmas)
        
            self.word_count = pd.DataFrame(
                word_count.most_common(), 
                columns = ['lemma', 'count']
                )
        
            self.exclude_list = self.word_count.loc[self.word_count['count']>= exclude_count]
            vocab = list(self.word_count.loc[self.word_count['count']<= exclude_count].lemma)
            self.vocab_list = [w.strip("""',.”"—-:;""") for w in vocab if w not in STOPS_LIST]
        
        # we now need to create a method to create paragraphs which will feed the doc reader 
            self.paras = []
            source = text_name
            doc_rows = self.reader.doc_rows(source)
            doc_rows = next(doc_rows) # set up dictionary iteration 
            # now iterate through doc_rows dictionary 
            for key in doc_rows:
                if (str(key).strip('<').strip('>').split('.')[-1]) == '0':
                    pass 
                elif (str(key).strip('<').strip('>').split('.')[-1]) == '1':
                    self.paras.append('\n' + str(doc_rows[key]))
                else:
                    self.paras.append(str(doc_rows[key]))

        
        # prepare formatting for Latin Library texts
        if self.corpus_name == 'latin_text_latin_library': 
            # change source to LL url syntax 
            url = 'http://thelatinlibrary.com/' + self.source.replace('txt', 'shtml')
            
            content = requests.get(url).content
            self.soup = BeautifulSoup(content, 'html.parser')
            self.paras = self.soup.body.get_text().replace('\n\n', '\n')
            # we need to read in the LL text and parse 
            
    def format_latin_library_text(soup, n: int=20): # not used 
        """
        Formats texts from the Latin Library. 
        :param soup: BeautifulSoup object instantiated 
        """
        if self.soup is not None: 
            for p in self.paras: 
                lines = para.split('\n')
                for line in range(0, len(lines), n):
                    chunk = lines[line:line + n]
            
        
    
    def clean_paragraph(self, ls): # not used 

        out = f""
        for i in range(len(ls) - 1):  
            if ls[i + 1] in punc: 
                out += ''.join([ls[i], ls[i + 1].strip('-')]) + ' '
            elif ls[i] not in punc:
                out += f'{ls[i]} '
            else: 
                pass
        return out.rstrip(' ')
         

    
    def parse_paragraph(self, paragraph: 'str') -> 'str': # not used 
        """Function to take paragraph of a Latin text and return a dictionary including definitions (but not citations). 
        The goal is to use this function to create short entries for a paragraph of a text. We can then use the paragraphs to build our Pharr formatted document. 
    
    :param paragraph: paragraph of parsed text
    """
    # TODO: refactor to reflect changes in notebook code 
    
        in_list = paragraph.split(' ')
        
        # filter list 
        
        in_list = [w for w in in_list if w not in STOPS_LIST]
        
        out_str = '' 
        path = 'lexica/Lewis_Short_XML/lat.ls.perseus-eng1.xml'
        tree = ET.parse(path)
        entries = tree.xpath('//entryFree')
        endings = ''
        gender = ""
        out_list = []

        for word in tqdm(in_list): 
            lemma = lemmatizer.lemmatize([word])[0][1]

            for entry in entries:
                senses = []
                if entry.get('key') == lemma:
                    if entry.find('itype') is not None: 
                        endings = f'{lemma} {entry.find("itype").text}'
                    if entry.find('gen') is not None: 
                        gender = entry.find('gen').text
                    for sense in entry.findall('sense')[:4]:
                        # print(sense.get('level'))
                        if sense.get("level") in ['1', '2']:
                            for tr in sense.findall('hi')[1:3]:
                                senses.append(tr.text)
                    out_string = f"""{endings} {gender}: {'; '.join(senses).strip('., ')}"""
                    out_list.append(out_string)
        return '\n'.join(out_list)
    
    def create_document(self,
                        paras: list=None,
                        output_format: '`poetry` or `prose`'='prose',
                        vocabulary_list = None,
                        title = '') -> ".docx file": 
        """
        Creates a .docx file with Latin text and vocabulary. 
        :param paras: paragraphs of text to include in the final document. If none supplied, will use paras from `get_text` method. 
        :param output_format: specifies the format of the final Word doc. If `poetry,` chunks will be ~20 lines. `poetry` setting splits the text by newline character `\n` If `prose` chunks will be of 150 words, splitting the text by spaces ` `. 
        :vocabulary list: special filtering list. If none supplied, words that appear 10 or more times in a text will be filtered out of facing vocabulary. 
        """
        doc = Document()
        count = 0
        if paras == None: 
            paras = '\n'.join(self.paras)
            print('generating document from entire text...')
        if output_format == 'poetry':
            line_size = 20
            self.paras = '\n'.join(self.paras)
            lines = paras.split('\n')
            joiner = '\n'
        elif output_format == 'prose': 
            line_size = 101
            self.paras = '\n'.join(self.paras)
            lines = self.paras.split(' ')
            joiner = ' '
        else:
            raise Exception("Format error. Spcify poetry or prose format")
        # if we have a Latin Library text, we grab the formatted text from the website
        if vocabulary_list == None: 
            vocabulary_list = self.vocab_list

        # we originally had this as dependent on the corpus name, let's try it without 
        # if corpus_name == 'tesserae':

        
        # if self.corpus_name == 'latin_text_latin_library': 
            # we have the paragraph list in self.paras 
            # break chunks 
            
        chunk_count = 0
        for chunk in tqdm(range(0, len(lines), line_size)):
            chunk_count +=1
            text_chunk = joiner.join(lines[chunk:chunk + line_size])
            # print("chunk_count: ", chunk_count, text_chunk)


            create_page(text_chunk = text_chunk,
                        vocabulary_list = vocabulary_list,
                        doc = doc,
                        save = False)
        if output_format == 'prose':
            line_size 
                
        doc.save(f'{title}.docx')